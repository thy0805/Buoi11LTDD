import os
import time
import shutil
import tkinter as tk
from tkinter import messagebox
import pyautogui
import docx

class RegionSelector:
    def __init__(self, parent):
        self.root = tk.Toplevel(parent)
        self.root.attributes("-alpha", 0.3)
        self.root.attributes("-fullscreen", True)
        self.root.attributes("-topmost", True)
        self.root.config(cursor="cross")
        
        self.canvas = tk.Canvas(self.root, cursor="cross", bg="grey")
        self.canvas.pack(fill="both", expand=True)
        
        self.canvas.bind("<ButtonPress-1>", self.on_press)
        self.canvas.bind("<B1-Motion>", self.on_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_release)
        
        self.start_x = None
        self.start_y = None
        self.rect = None
        self.region = None
        
    def on_press(self, event):
        self.start_x = event.x
        self.start_y = event.y
        self.rect = self.canvas.create_rectangle(self.start_x, self.start_y, 1, 1, outline="red", width=2)
        
    def on_drag(self, event):
        cur_x, cur_y = event.x, event.y
        self.canvas.coords(self.rect, self.start_x, self.start_y, cur_x, cur_y)
        
    def on_release(self, event):
        end_x, end_y = event.x, event.y
        x1 = min(self.start_x, end_x)
        y1 = min(self.start_y, end_y)
        x2 = max(self.start_x, end_x)
        y2 = max(self.start_y, end_y)
        w = x2 - x1
        h = y2 - y1
        if w > 10 and h > 10:
            self.region = (x1, y1, w, h)
        self.root.destroy()
        
    def select(self):
        self.root.grab_set()
        self.root.wait_window()
        return self.region

class CaptureApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("HUIT Capture Menu")
        self.root.geometry("460x280")
        self.root.config(bg="#F5F7FA")
        self.root.resizable(False, False)
        
        title_label = tk.Label(
            self.root,
            text="HUIT CAPTURE UTILITY",
            font=("Arial", 14, "bold"),
            bg="#F5F7FA",
            fg="#1E88E5"
        )
        title_label.pack(pady=16)
        
        select_label = tk.Label(
            self.root,
            text="Chọn bài tập để chụp:",
            font=("Arial", 10),
            bg="#F5F7FA",
            fg="#37474F"
        )
        select_label.pack(pady=4)
        
        self.selected_bai = tk.StringVar(self.root)
        self.selected_bai.set("Bài 1")
        
        self.option_menu = tk.OptionMenu(
            self.root,
            self.selected_bai,
            "Bài 1", "Bài 2", "Bài 3", "Bài 4"
        )
        self.option_menu.config(width=15, font=("Arial", 10))
        self.option_menu.pack(pady=8)
        
        self.btn_capture = tk.Button(
            self.root,
            text="Chụp & Chèn Vào Word",
            font=("Arial", 11, "bold"),
            bg="#1E88E5",
            fg="white",
            relief="flat",
            command=self.start_capture
        )
        self.btn_capture.pack(pady=16, ipadx=10, ipady=4)
        
    def start_capture(self):
        try:
            bai_str = self.selected_bai.get()
            bai_num = int(bai_str.split(" ")[1])
            
            folders = {
                1: "Buoi11Tailop",
                2: "Buoi11Tailop",
                3: "Buoi11Tailop",
                4: "Buoi11Venha"
            }
            
            source_file = f"c:\\hoc tap\\LTDD\\Buoi11\\{folders[bai_num]}\\bai{bai_num}\\lib\\main.dart"
            
            if not os.path.exists(source_file):
                messagebox.showerror("Lỗi", f"Không tìm thấy file source: {source_file}")
                return
                
            self.root.withdraw()
            
            selector = RegionSelector(self.root)
            region = selector.select()
            
            if not region:
                return
                
            time.sleep(5)
            
            with open(source_file, "r", encoding="utf-8") as f:
                lines = f.readlines()
            line_count = len(lines)
            num_scrolls = (line_count // 26) + 1
            
            temp_dir = f"C:\\Users\\thy\\.gemini\\antigravity-ide\\brain\\2e5688ae-24b4-4fcb-8040-e8006babad60\\scratch\\gui_temp_bai{bai_num}"
            os.makedirs(temp_dir, exist_ok=True)
            
            code_image_paths = []
            for i in range(num_scrolls):
                img_path = os.path.join(temp_dir, f"code_region_{i}.png")
                screenshot = pyautogui.screenshot(region=region)
                screenshot.save(img_path)
                code_image_paths.append(img_path)
                
                for _ in range(26):
                    pyautogui.press("down")
                    time.sleep(0.02)
                    
                time.sleep(0.8)
                
            demo_image_paths = []
            capture_demo = messagebox.askyesno(
                "Chụp Giao Diện",
                "Đã chụp xong Code. Thy có muốn chụp ảnh giao diện (UI Demo) cho bài này không?",
                parent=self.root
            )
            
            if capture_demo:
                demo_idx = 0
                while True:
                    messagebox.showinfo(
                        "Hướng dẫn",
                        "Chuẩn bị giao diện trên máy ảo, bấm OK rồi kéo chọn vùng màn hình muốn chụp nha Thy!",
                        parent=self.root
                    )
                    
                    demo_selector = RegionSelector(self.root)
                    demo_region = demo_selector.select()
                    
                    if not demo_region:
                        break
                        
                    time.sleep(1.5)
                    demo_path = os.path.join(temp_dir, f"demo_region_{demo_idx}.png")
                    screenshot_demo = pyautogui.screenshot(region=demo_region)
                    screenshot_demo.save(demo_path)
                    demo_image_paths.append(demo_path)
                    demo_idx += 1
                    
                    ask_more = messagebox.askyesno(
                        "Chụp Tiếp",
                        f"Đã chụp thành công {demo_idx} ảnh giao diện. Thy có muốn chụp thêm ảnh giao diện khác nữa không?",
                        parent=self.root
                    )
                    if not ask_more:
                        break
                        
            output_docx = r"c:\hoc tap\LTDD\Buoi11\buoi11.docx"
            
            if os.path.exists(output_docx):
                try:
                    doc = docx.Document(output_docx)
                except Exception:
                    doc = docx.Document()
            else:
                doc = docx.Document()
                
            sections = doc.sections
            for section in sections:
                section.top_margin = docx.shared.Inches(0.5)
                section.bottom_margin = docx.shared.Inches(0.5)
                section.left_margin = docx.shared.Inches(0.5)
                section.right_margin = docx.shared.Inches(0.5)
                
            idx_curr = -1
            idx_next = -1
            
            for idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip().upper()
                if text == f"BÀI {bai_num}":
                    idx_curr = idx
                elif text == f"BÀI {bai_num + 1}":
                    idx_next = idx
                    break
                    
            body_element = doc.element.body
            
            if idx_curr != -1:
                if idx_next != -1:
                    p_next = doc.paragraphs[idx_next]
                    for idx in range(idx_next - 1, idx_curr - 1, -1):
                        p = doc.paragraphs[idx]
                        body_element.remove(p._element)
                    target_p = p_next
                else:
                    for idx in range(len(doc.paragraphs) - 1, idx_curr - 1, -1):
                        p = doc.paragraphs[idx]
                        body_element.remove(p._element)
                    target_p = None
            else:
                target_p = None
                for idx, paragraph in enumerate(doc.paragraphs):
                    text = paragraph.text.strip().upper()
                    if text.startswith("BÀI "):
                        try:
                            num = int(text.split(" ")[1])
                            if num > bai_num:
                                target_p = paragraph
                                break
                        except Exception:
                            pass
                            
            if target_p is not None:
                h = target_p.insert_paragraph_before()
                h.paragraph_format.space_before = docx.shared.Pt(24)
                h.paragraph_format.space_after = docx.shared.Pt(12)
                run_text = h.add_run(f"BÀI {bai_num}")
                run_text.font.name = "Arial"
                run_text.font.size = docx.shared.Pt(13)
                run_text.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                run_text.font.bold = True
                
                for path in code_image_paths:
                    p_img = target_p.insert_paragraph_before()
                    p_img.paragraph_format.space_before = docx.shared.Pt(6)
                    p_img.paragraph_format.space_after = docx.shared.Pt(6)
                    run_img = p_img.add_run()
                    run_img.add_picture(path, width=docx.shared.Inches(7.0))
                    
                if demo_image_paths:
                    p_title = target_p.insert_paragraph_before()
                    p_title.paragraph_format.space_before = docx.shared.Pt(12)
                    p_title.paragraph_format.space_after = docx.shared.Pt(6)
                    run_title = p_title.add_run("ẢNH CHỤP GIAO DIỆN DEMO")
                    run_title.font.name = "Arial"
                    run_title.font.size = docx.shared.Pt(11)
                    run_title.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                    run_title.font.bold = True
                    
                    for path in demo_image_paths:
                        p_img = target_p.insert_paragraph_before()
                        p_img.paragraph_format.space_before = docx.shared.Pt(6)
                        p_img.paragraph_format.space_after = docx.shared.Pt(6)
                        run_img = p_img.add_run()
                        run_img.add_picture(path, width=docx.shared.Inches(7.0))
            else:
                heading = doc.add_heading(level=1)
                heading.paragraph_format.space_before = docx.shared.Pt(24)
                heading.paragraph_format.space_after = docx.shared.Pt(12)
                run_text = heading.add_run(f"BÀI {bai_num}")
                run_text.font.name = "Arial"
                run_text.font.size = docx.shared.Pt(13)
                run_text.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                run_text.font.bold = True
                
                for path in code_image_paths:
                    doc.add_picture(path, width=docx.shared.Inches(7.0))
                    doc.add_paragraph()
                    
                if demo_image_paths:
                    p_title = doc.add_paragraph()
                    p_title.paragraph_format.space_before = docx.shared.Pt(12)
                    p_title.paragraph_format.space_after = docx.shared.Pt(6)
                    run_title = p_title.add_run("ẢNH CHỤP GIAO DIỆN DEMO")
                    run_title.font.name = "Arial"
                    run_title.font.size = docx.shared.Pt(11)
                    run_title.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
                    run_title.font.bold = True
                    
                    for path in demo_image_paths:
                        doc.add_picture(path, width=docx.shared.Inches(7.0))
                        doc.add_paragraph()
                        
            doc.save(output_docx)
            shutil.rmtree(temp_dir)
            
            self.root.deiconify()
            messagebox.showinfo("Thành công", f"Đã chụp và chèn thành công BÀI {bai_num} vào đúng vị trí!", parent=self.root)
        except Exception as e:
            self.root.deiconify()
            messagebox.showerror("Lỗi", f"Có lỗi xảy ra: {str(e)}\n\n(Nếu là lỗi Permission denied, hãy đóng file Word!)", parent=self.root)
        finally:
            self.root.deiconify()
        
    def start(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = CaptureApp()
    app.start()
